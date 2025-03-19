"""
WordDocumentAdapter: Adapter for processing Microsoft Word (.docx, .doc) files.
"""

import os
import logging
import tempfile
import time
from typing import Dict, List, Optional, Tuple, Union, Any
import traceback
from datetime import datetime
import hashlib

from models_app.vision.document.adapters.document_format_adapter import DocumentFormatAdapter
from models_app.vision.document.utils.core.processing_metadata_context import ProcessingMetadataContext
from models_app.vision.document.utils.error_handling.errors import (
    DocumentProcessingError,
    DocumentValidationError
)
from models_app.vision.document.utils.error_handling.handlers import (
    handle_document_errors,
    measure_processing_time
)
from models_app.vision.document.utils.core.next_layer_interface import ProcessingEventType

logger = logging.getLogger(__name__)

class WordDocumentAdapter(DocumentFormatAdapter):
    """
    Adapter for processing Microsoft Word (.docx, .doc) files.
    Uses python-docx library to extract text, structure, and metadata.
    
    Features:
    - Extracts full text content with formatting
    - Preserves document structure (headings, paragraphs, lists)
    - Extracts tables and converts them to structured data
    - Identifies images and embedded objects
    - Retrieves document metadata (author, title, etc.)
    """
    
    # Class-level constants
    VERSION = "1.0.0"
    CAPABILITIES = {
        "text_extraction": 0.9,
        "structure_preservation": 0.8,
        "metadata_extraction": 0.8,
        "table_extraction": 0.7,
        "image_extraction": 0.5
    }
    SUPPORTED_FORMATS = [".docx", ".doc", ".rtf", ".odt"]
    PRIORITY = 80
    
    def __init__(self, config: Dict[str, Any] = None):
        """
        Initialize the Word document adapter.
        
        Args:
            config: Configuration dictionary with adapter-specific settings.
        """
        super().__init__(config)
        self.docx_module = None
        self.document = None
    
    def _initialize_adapter_components(self) -> None:
        """Initialize adapter-specific components."""
        try:
            # Attempt to import python-docx
            import docx
            from docx.document import Document as DocxDocument
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import _Cell, Table
            from docx.text.paragraph import Paragraph
            
            self.docx_module = docx
            logger.info("Successfully initialized Word document processing components")
            
            # Additional components for .doc files (optional)
            try:
                import win32com.client
                self.has_win32com = True
            except ImportError:
                self.has_win32com = False
                logger.info("win32com not available - .doc file handling will be limited")
            
        except ImportError as e:
            logger.warning(f"Could not import Word processing libraries: {str(e)}. Using fallback mode.")
            self._initialize_fallback_components()
    
    def _initialize_fallback_components(self) -> None:
        """Initialize fallback components when main libraries are unavailable."""
        self.can_process_docx = False
        logger.info("Initialized fallback mode - document extraction will be limited to text only")
    
    @handle_document_errors
    @measure_processing_time
    def _process_file(self, file_path: str, options: Dict[str, Any], 
                     metadata_context: Optional[ProcessingMetadataContext] = None) -> Dict[str, Any]:
        """
        Process a Word document file and extract information.
        
        Args:
            file_path: Path to the Word document file
            options: Processing options
            metadata_context: Context for tracking metadata during processing
            
        Returns:
            Dictionary with extracted information
        """
        if metadata_context:
            metadata_context.start_timing("word_processing")
            metadata_context.record_adapter_processing(
                adapter=self.__class__.__name__,
                file_path=file_path,
                capabilities=self.CAPABILITIES
            )
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Record processing start for performance tracking
        start_time = time.time()
        
        # Emit start event
        self.next_layer.emit_simple_event(
            event_type=ProcessingEventType.PROCESSING_PHASE_START,
            document_id=file_path,
            data={
                "adapter": self.__class__.__name__,
                "file_extension": file_extension,
                "phase": "word_processing"
            }
        )
        
        try:
            # Process based on file extension
            if file_extension == '.docx':
                result = self._process_docx_file(file_path, options, metadata_context)
            elif file_extension == '.doc':
                result = self._process_doc_file(file_path, options, metadata_context)
            elif file_extension in ['.odt', '.rtf']:
                result = self._process_other_document(file_path, options, metadata_context)
            else:
                raise DocumentValidationError(
                    f"Unsupported file format: {file_extension}",
                    document_path=file_path
                )
            
            # Calculate processing time for performance tracking
            processing_time = time.time() - start_time
            
            # Add processor info to result
            result = {
                **result,
                "processor": {
                    "name": self.__class__.__name__,
                    "version": self.VERSION,
                    "processing_time": processing_time,
                    "timestamp": datetime.now().isoformat()
                }
            }
            
            # Emit performance metric
            self.next_layer.emit_simple_event(
                event_type=ProcessingEventType.PERFORMANCE_METRIC,
                document_id=file_path,
                data={
                    "adapter": self.__class__.__name__,
                    "processing_time": processing_time,
                    "success": True,
                    "file_extension": file_extension
                }
            )
            
            # Emit completion event
            self.next_layer.emit_simple_event(
                event_type=ProcessingEventType.PROCESSING_PHASE_END,
                document_id=file_path,
                data={
                    "adapter": self.__class__.__name__,
                    "phase": "word_processing",
                    "success": True,
                    "processing_time": processing_time
                }
            )
            
            if metadata_context:
                metadata_context.end_timing("word_processing")
                
            return result
            
        except Exception as e:
            # Calculate error processing time
            error_time = time.time() - start_time
            
            # Emit error event
            self.next_layer.emit_simple_event(
                event_type=ProcessingEventType.ERROR_OCCURRED,
                document_id=file_path,
                data={
                    "adapter": self.__class__.__name__,
                    "phase": "word_processing",
                    "error": str(e),
                    "error_type": type(e).__name__,
                    "processing_time": error_time
                }
            )
            
            if metadata_context:
                metadata_context.record_error(
                    component=self.__class__.__name__,
                    message=f"Error processing Word document: {str(e)}",
                    error_type=type(e).__name__
                )
                metadata_context.end_timing("word_processing")
            
            logger.error(f"Error processing Word document: {str(e)}")
            logger.debug(traceback.format_exc())
            
            raise DocumentProcessingError(f"Error processing Word document: {str(e)}")
    
    def _process_docx_file(self, file_path: str, options: Dict[str, Any],
                          metadata_context: Optional[ProcessingMetadataContext] = None) -> Dict[str, Any]:
        """
        Process a .docx file using python-docx.
        
        Args:
            file_path: Path to the .docx file
            options: Processing options
            metadata_context: Context for tracking metadata
            
        Returns:
            Dictionary with extracted content and structure
        """
        if not self.docx_module:
            raise DocumentProcessingError(
                "Word processing libraries not available. Cannot process .docx file."
            )
        
        try:
            # Load document
            self.document = self.docx_module.Document(file_path)
            
            # Extract metadata
            metadata = self._extract_document_properties()
            
            # Process document structure
            paragraphs_data = []
            tables_data = []
            images_data = []
            all_text = []
            
            # Process all paragraphs, tables, and other elements
            for element in self.document.element.body:
                if isinstance(element, self.docx_module.oxml.text.paragraph.CT_P):
                    # Process paragraph
                    paragraph = self.docx_module.text.paragraph.Paragraph(element, self.document)
                    paragraph_data = self._process_paragraph(paragraph)
                    paragraphs_data.append(paragraph_data)
                    all_text.append(paragraph_data["text"])
                    
                elif isinstance(element, self.docx_module.oxml.table.CT_Tbl):
                    # Process table
                    table = self.docx_module.table.Table(element, self.document)
                    table_data = self._process_table(table)
                    tables_data.append(table_data)
                    # Add table text to all text
                    table_text = "\n".join([cell for row in table_data["data"] for cell in row if cell])
                    all_text.append(table_text)
            
            # Extract images if available (basic implementation)
            try:
                for rel in self.document.part.rels.values():
                    if "image" in rel.target_ref:
                        image_data = {
                            "id": rel.rId,
                            "filename": os.path.basename(rel.target_ref),
                            "content_type": rel.target_part.content_type,
                            "size": len(rel.target_part.blob) if hasattr(rel.target_part, "blob") else 0
                        }
                        images_data.append(image_data)
            except Exception as image_err:
                logger.warning(f"Error extracting images: {str(image_err)}")
            
            # Get document styles
            styles_data = self._extract_styles()
            
            # Estimate page count
            page_count = self._estimate_page_count()
            
            # Prepare final result
            result = {
                "document_type": "docx",
                "file_path": file_path,
                "text": "\n\n".join(all_text),
                "metadata": metadata,
                "structure": {
                    "paragraphs": paragraphs_data,
                    "tables": tables_data,
                    "images": images_data,
                    "styles": styles_data,
                    "page_count": page_count
                },
                "processing_timestamp": datetime.now().isoformat()
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing .docx file {file_path}: {str(e)}")
            raise DocumentProcessingError(f"Failed to process .docx file: {str(e)}")
    
    def _process_doc_file(self, file_path: str, options: Dict[str, Any],
                         metadata_context: Optional[ProcessingMetadataContext] = None) -> Dict[str, Any]:
        """
        Process a .doc file.
        
        Uses win32com if available, otherwise attempts basic text extraction.
        
        Args:
            file_path: Path to the .doc file
            options: Processing options
            metadata_context: Context for tracking metadata
            
        Returns:
            Dictionary with extracted content
        """
        # Check if we have win32com for proper .doc processing
        if self.has_win32com:
            try:
                # Convert .doc to .docx using win32com
                import win32com.client
                import os
                
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                # Create temporary file for the conversion
                temp_file = self._create_temp_file(suffix=".docx")
                
                # Open and save as docx
                doc = word.Documents.Open(os.path.abspath(file_path))
                doc.SaveAs(temp_file, 16)  # 16 = docx format
                doc.Close()
                word.Quit()
                
                # Process the converted docx
                result = self._process_docx_file(temp_file, options, metadata_context)
                result["document_type"] = "doc"
                return result
                
            except Exception as e:
                logger.warning(f"Failed to convert .doc using win32com: {str(e)}. Falling back to basic extraction.")
        
        # Fallback to basic text extraction
        try:
            # Try using textract if available
            try:
                import textract
                text = textract.process(file_path).decode('utf-8')
            except ImportError:
                # Very basic fallback - just read file as binary and extract printable chars
                with open(file_path, 'rb') as f:
                    content = f.read()
                text = ''.join(chr(c) for c in content if 32 <= c < 127 or c in [9, 10, 13])
            
            # Create simplified result
            return {
                "document_type": "doc",
                "file_path": file_path,
                "text": text,
                "metadata": {
                    "filename": os.path.basename(file_path),
                    "file_size": os.path.getsize(file_path)
                },
                "structure": {
                    "paragraphs": [],  # Cannot extract structure
                    "tables": [],      # Cannot extract tables
                    "images": [],      # Cannot extract images
                    "styles": {},      # Cannot extract styles
                    "page_count": 0    # Unknown page count
                },
                "processing_timestamp": datetime.now().isoformat()
            }
        except Exception as e:
            logger.error(f"Error processing .doc file {file_path}: {str(e)}")
            raise DocumentProcessingError(f"Failed to process .doc file: {str(e)}")
    
    def _process_other_document(self, file_path: str, options: Dict[str, Any],
                              metadata_context: Optional[ProcessingMetadataContext] = None) -> Dict[str, Any]:
        """
        Process .rtf or .odt files.
        
        Uses textract if available, otherwise attempts very basic text extraction.
        
        Args:
            file_path: Path to the document file
            options: Processing options
            metadata_context: Context for tracking metadata
            
        Returns:
            Dictionary with basic extracted content
        """
        try:
            # Try using textract if available
            try:
                import textract
                text = textract.process(file_path).decode('utf-8')
            except ImportError:
                # Basic fallback - just read file and extract printable chars
                with open(file_path, 'rb') as f:
                    content = f.read()
                text = ''.join(chr(c) for c in content if 32 <= c < 127 or c in [9, 10, 13])
            
            # Create simplified result
            file_extension = os.path.splitext(file_path)[1].lower()[1:]  # Remove the dot
            return {
                "document_type": file_extension,
                "file_path": file_path,
                "text": text,
                "metadata": {
                    "filename": os.path.basename(file_path),
                    "file_size": os.path.getsize(file_path)
                },
                "structure": {
                    "paragraphs": [],  # Cannot extract structure
                    "tables": [],      # Cannot extract tables
                    "images": [],      # Cannot extract images
                    "styles": {},      # Cannot extract styles
                    "page_count": 0    # Unknown page count
                },
                "processing_timestamp": datetime.now().isoformat()
            }
        except Exception as e:
            logger.error(f"Error processing {file_path}: {str(e)}")
            raise DocumentProcessingError(f"Failed to process {file_path}: {str(e)}")
    
    def _process_paragraph(self, paragraph) -> Dict[str, Any]:
        """
        Process a paragraph and extract its text and properties.
        
        Args:
            paragraph: Document paragraph object
            
        Returns:
            Dictionary with paragraph data
        """
        try:
            # Get paragraph text
            text = paragraph.text
            
            # Get paragraph style
            style = paragraph.style.name if hasattr(paragraph, "style") and paragraph.style else "Normal"
            
            # Get alignment
            alignment = "left"  # Default
            if hasattr(paragraph, "paragraph_format") and paragraph.paragraph_format:
                if paragraph.paragraph_format.alignment:
                    alignments = {
                        0: "left",
                        1: "center",
                        2: "right",
                        3: "justify"
                    }
                    alignment = alignments.get(paragraph.paragraph_format.alignment, "left")
            
            # Determine if heading
            is_heading = style.startswith("Heading")
            heading_level = 0
            if is_heading:
                try:
                    heading_level = int(style.replace("Heading", "").strip())
                except ValueError:
                    heading_level = 0
            
            # Check if list item
            is_list_item = False
            if paragraph._p.pPr and paragraph._p.pPr.numPr:
                is_list_item = True
            
            return {
                "text": text,
                "style": style,
                "alignment": alignment,
                "is_heading": is_heading,
                "heading_level": heading_level,
                "is_list_item": is_list_item,
                "length": len(text)
            }
        except Exception as e:
            logger.warning(f"Error processing paragraph: {str(e)}")
            return {
                "text": paragraph.text if hasattr(paragraph, "text") else "",
                "error": str(e)
            }
    
    def _process_table(self, table) -> Dict[str, Any]:
        """
        Process a table and extract its structure and content.
        
        Args:
            table: Document table object
            
        Returns:
            Dictionary with table data
        """
        try:
            rows = []
            headers = []
            
            # Process all rows
            for i, row in enumerate(table.rows):
                cells = []
                for cell in row.cells:
                    # Get cell text (join paragraphs)
                    cell_text = "\n".join(p.text for p in cell.paragraphs)
                    cells.append(cell_text)
                
                # Store first row as headers
                if i == 0:
                    headers = cells
                
                rows.append(cells)
            
            return {
                "rows": len(rows),
                "columns": len(headers),
                "headers": headers,
                "data": rows,
                "has_header": True  # Assume first row is header by default
            }
        except Exception as e:
            logger.warning(f"Error processing table: {str(e)}")
            return {
                "rows": 0,
                "columns": 0,
                "headers": [],
                "data": [],
                "error": str(e)
            }
    
    def _extract_styles(self) -> Dict[str, Any]:
        """
        Extract document styles information.
        
        Returns:
            Dictionary with styles data
        """
        if not self.document:
            return {}
            
        try:
            styles = {}
            for style in self.document.styles:
                if hasattr(style, "name"):
                    styles[style.name] = {
                        "type": style.type if hasattr(style, "type") else "unknown",
                        "is_builtin": style.builtin if hasattr(style, "builtin") else False
                    }
            return styles
        except Exception as e:
            logger.warning(f"Error extracting styles: {str(e)}")
            return {}
    
    def _extract_document_properties(self) -> Dict[str, Any]:
        """
        Extract document properties/metadata.
        
        Returns:
            Dictionary with document properties
        """
        if not self.document:
            return {}
            
        try:
            core_properties = self.document.core_properties
            
            metadata = {
                "title": core_properties.title if hasattr(core_properties, "title") else "",
                "author": core_properties.author if hasattr(core_properties, "author") else "",
                "subject": core_properties.subject if hasattr(core_properties, "subject") else "",
                "keywords": core_properties.keywords if hasattr(core_properties, "keywords") else "",
                "comments": core_properties.comments if hasattr(core_properties, "comments") else "",
                "category": core_properties.category if hasattr(core_properties, "category") else "",
                "created": str(core_properties.created) if hasattr(core_properties, "created") else "",
                "modified": str(core_properties.modified) if hasattr(core_properties, "modified") else "",
                "last_modified_by": core_properties.last_modified_by if hasattr(core_properties, "last_modified_by") else "",
                "revision": core_properties.revision if hasattr(core_properties, "revision") else 0,
                "content_status": core_properties.content_status if hasattr(core_properties, "content_status") else ""
            }
            
            return metadata
        except Exception as e:
            logger.warning(f"Error extracting document properties: {str(e)}")
            return {}
    
    def _estimate_page_count(self) -> int:
        """
        Estimate number of pages in the document.
        This is an approximation as python-docx doesn't provide actual page info.
        
        Returns:
            Estimated page count
        """
        if not self.document:
            return 0
            
        try:
            # Rough estimate based on paragraphs and tables
            paragraph_count = len(self.document.paragraphs)
            table_count = len(self.document.tables)
            
            # Very rough estimate: ~40 paragraphs per page, tables count as 5 paragraphs
            estimated_pages = (paragraph_count + (table_count * 5)) / 40
            return max(1, int(estimated_pages))
        except Exception as e:
            logger.warning(f"Error estimating page count: {str(e)}")
            return 1

    def validate_format(self, document_path: str) -> bool:
        """Validate if the document format is supported."""
        _, ext = os.path.splitext(document_path)
        return ext.lower() in self.SUPPORTED_FORMATS
    
    def extract_structure(self, document_path: str, metadata_context: Optional[ProcessingMetadataContext] = None) -> Dict[str, Any]:
        """Extract document structure (headings, paragraphs, etc.)."""
        # Implementation that extracts structural elements
        structure = {
            "type": "word_document",
            "elements": []
        }
        
        # Extract headings, paragraphs, tables, etc.
        if self.docx_module and document_path:
            try:
                self.document = self.docx_module.Document(document_path)
                
                # Extract paragraphs
                paragraphs = []
                for p in self.document.paragraphs:
                    if p.text.strip():
                        paragraphs.append({
                            "text": p.text,
                            "style": p.style.name if hasattr(p, "style") and p.style else "Normal"
                        })
                
                structure["elements"] = paragraphs
                
                # Extract tables
                tables = []
                for table in self.document.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables.append({
                        "rows": len(table.rows),
                        "columns": len(table.rows[0].cells) if table.rows else 0,
                        "data": table_data
                    })
                
                structure["tables"] = tables
                
            except Exception as e:
                logger.error(f"Error extracting structure: {str(e)}")
        
        return structure

# Register this adapter
DocumentFormatAdapter.register_adapter("word", WordDocumentAdapter)